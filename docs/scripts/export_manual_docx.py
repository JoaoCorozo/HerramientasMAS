"""
Exporta MANUAL_USUARIO_PLATAFORMA_BEX.md a Word (.docx) con capturas.
Uso: py docs/scripts/export_manual_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

DOCS = Path(__file__).resolve().parents[1]
MD_FILE = DOCS / "MANUAL_USUARIO_PLATAFORMA_BEX.md"
OUT_DOCX = DOCS / "Manual_Usuario_Plataforma_BEX.docx"
CAPTURAS = DOCS / "capturas"


def add_rich_text(paragraph, text: str, base_bold: bool = False):
    """Inserta texto con **negrita** y `código`."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
            if base_bold:
                run.bold = True


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip("|").split("|")]


def is_table_sep(line: str) -> bool:
    cells = parse_table_row(line)
    return all(set(c) <= {"-", ":", " "} for c in cells)


def export():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Manual de usuario\nPlataforma de Herramientas BEX")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(15, 76, 129)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Guía para usuarios · Mayo 2026\nUso interno — BEX")
    s.font.size = Pt(12)
    s.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_page_break()

    lines = MD_FILE.read_text(encoding="utf-8").splitlines()
    i = 0
    in_table = False
    table_rows: list[list[str]] = []
    list_items: list[str] = []
    list_ordered = False

    def flush_list():
        nonlocal list_items, list_ordered
        for idx, item in enumerate(list_items, 1):
            p = doc.add_paragraph(style="List Number" if list_ordered else "List Bullet")
            add_rich_text(p, item)
        list_items = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = "Table Grid"
        for ri, row in enumerate(table_rows):
            for ci, cell in enumerate(row):
                cell_p = table.rows[ri].cells[ci].paragraphs[0]
                add_rich_text(cell_p, cell, base_bold=(ri == 0))
        table_rows = []
        in_table = False
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        if line.startswith("![") and "](" in line:
            flush_list()
            flush_table()
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
            if m:
                cap, rel = m.group(1), m.group(2)
                img_path = DOCS / rel if not Path(rel).is_absolute() else Path(rel)
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Inches(6.0))
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_p.add_run(cap)
                    cap_run.italic = True
                    cap_run.font.size = Pt(9)
                    cap_run.font.color.rgb = RGBColor(100, 100, 100)
                else:
                    doc.add_paragraph(f"[Imagen no encontrada: {rel}]")
            continue

        if line.startswith("# "):
            flush_list()
            flush_table()
            doc.add_heading(line[2:], level=0)
            continue
        if line.startswith("## "):
            flush_list()
            flush_table()
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            flush_list()
            flush_table()
            doc.add_heading(line[4:], level=2)
            continue
        if line == "---":
            flush_list()
            flush_table()
            continue
        if line.startswith("|"):
            flush_list()
            if is_table_sep(line):
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(parse_table_row(line))
            continue
        else:
            if in_table:
                flush_table()

        if line.startswith("> "):
            flush_list()
            p = doc.add_paragraph(style="Intense Quote")
            add_rich_text(p, line[2:])
            continue

        if re.match(r"^\d+\.\s", line):
            if list_items and not list_ordered:
                flush_list()
            list_ordered = True
            list_items.append(re.sub(r"^\d+\.\s", "", line))
            continue

        if line.startswith("- "):
            if list_items and list_ordered:
                flush_list()
            list_ordered = False
            list_items.append(line[2:])
            continue

        if list_items and line.strip():
            flush_list()

        if not line.strip():
            continue

        # Saltar bloque de TOC con enlaces markdown
        if line.strip().startswith("1. [") and "](#" in line:
            continue

        p = doc.add_paragraph()
        add_rich_text(p, line)

    flush_list()
    flush_table()

    doc.save(str(OUT_DOCX))
    print(f"Word generado: {OUT_DOCX}")
    print(f"Tamaño: {OUT_DOCX.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    export()
